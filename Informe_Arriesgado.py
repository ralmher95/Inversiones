"""
Informe Cartera Arriesgada - Word (Monte Carlo, sin scipy)
============================================================
"""

import pandas as pd
import numpy as np
import yfinance as yf
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

ACTIVOS = {
    "CSPX.L": "S&P 500",
    "QDVE.DE": "Info Tech Sector",
    "IUSN.DE": "World Small Cap",
    "IS3N.DE": "Emerging Markets",
    "BTCE.DE": "Bitcoin ETC",
}
TICKERS = list(ACTIVOS.keys())
TASA_RF = 0.0315
N_DIAS = 252
PERIODO = "3y"
MC_SIMULACIONES = 50000

print("⬇ Descargando...")
raw = yf.download(TICKERS, period=PERIODO, auto_adjust=True, progress=False)

if "Close" not in raw.columns:
    data = raw.dropna(axis=1, how='all')
else:
    data = raw["Close"].dropna(axis=1, how='all')

min_dias = 100
activos_ok = [t for t in data.columns if data[t].notna().sum() >= min_dias]
if len(activos_ok) < 2:
    raise ValueError("Se necesitan al menos 2 activos con datos.")
data = data[activos_ok]
nombres_ok = [ACTIVOS.get(t, t) for t in activos_ok]
returns = data.pct_change().dropna()
n = len(activos_ok)
print(f"✅ Activos: {activos_ok}\n")

ret_anual = returns.mean() * N_DIAS
vol_anual = returns.std() * np.sqrt(N_DIAS)
sharpe = (ret_anual - TASA_RF) / vol_anual
max_dd = ((data / data.cummax()) - 1).min()

activos_rows = []
for t in activos_ok:
    activos_rows.append([
        f"{ACTIVOS.get(t, t)} ({t})",
        f"{ret_anual[t]*100:.2f}",
        f"{vol_anual[t]*100:.2f}",
        f"{max_dd[t]*100:.2f}",
        f"{sharpe[t]:.3f}"
    ])

corr = returns.corr()
corr_named = corr.copy()
corr_named.columns = nombres_ok
corr_named.index = nombres_ok

def stats_cartera(weights):
    weights = np.array(weights)
    port_returns = returns.dot(weights)
    ret = port_returns.mean() * N_DIAS
    vol = port_returns.std() * np.sqrt(N_DIAS)
    if vol == 0: return np.nan, np.nan, np.nan
    sr = (ret - TASA_RF) / vol
    return ret, vol, sr

print(f"🎲 Monte Carlo {MC_SIMULACIONES:,} carteras...")
np.random.seed(42)
mc_ret = np.zeros(MC_SIMULACIONES)
mc_vol = np.zeros(MC_SIMULACIONES)
mc_sr  = np.zeros(MC_SIMULACIONES)
pesos_mc = np.zeros((MC_SIMULACIONES, n))

for i in range(MC_SIMULACIONES):
    w = np.random.random(n)
    w /= w.sum()
    pesos_mc[i] = w
    r, v, s = stats_cartera(w)
    mc_ret[i] = r
    mc_vol[i] = v
    mc_sr[i]  = s

idx_sharpe = np.nanargmax(mc_sr)
idx_vol = np.nanargmin(mc_vol)
best_w_sharpe = pesos_mc[idx_sharpe]
best_w_vol = pesos_mc[idx_vol]
res_s = stats_cartera(best_w_sharpe)
res_v = stats_cartera(best_w_vol)

pesos_s = {nombres_ok[i]: best_w_sharpe[i]*100 for i in range(n) if best_w_sharpe[i] > 0.005}
pesos_v = {nombres_ok[i]: best_w_vol[i]*100 for i in range(n) if best_w_vol[i] > 0.005}

# ── Documento Word ────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

title = doc.add_heading('INFORME FINANCIERO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Cartera Arriesgada Global (5 ETFs + Bitcoin)')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x33, 0x55, 0x77)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")} | Periodo: {PERIODO} | RF BCE: {TASA_RF*100:.2f}%').font.size = Pt(9)
doc.add_paragraph()

def add_table(headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for row in data:
        rc = table.add_row().cells
        for i, item in enumerate(row):
            rc[i].text = str(item)
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

doc.add_heading('1. Resumen Ejecutivo', level=1)
doc.add_paragraph(
    f'Datos desde {returns.index[0].date()} hasta {returns.index[-1].date()}. '
    f'Optimización Monte Carlo con {MC_SIMULACIONES:,} carteras.'
)
add_table(
    ['Cartera', 'Retorno (%)', 'Volatilidad (%)', 'Sharpe'],
    [
        ['Máximo Sharpe', f"{res_s[0]*100:.1f}", f"{res_s[1]*100:.1f}", f"{res_s[2]:.3f}"],
        ['Mínima Volatilidad', f"{res_v[0]*100:.1f}", f"{res_v[1]*100:.1f}", f"{res_v[2]:.3f}"],
        ['Tasa libre riesgo', f"{TASA_RF*100:.2f}", "0.00", '—']
    ]
)

doc.add_heading('2. Composición Orientativa (Pesos Típicos)', level=1)
doc.add_paragraph('Referencia para perfil arriesgado. Los pesos reales se muestran en la sección 5.')
add_table(
    ['Activo', 'Rol'],
    [
        ['S&P 500 (CSPX.L)', 'Núcleo renta variable USA'],
        ['Info Tech Sector (QDVE.DE)', 'Alto crecimiento tecnológico'],
        ['World Small Cap (IUSN.DE)', 'Small caps globales'],
        ['Emerging Markets (IS3N.DE)', 'Renta variable emergente'],
        ['Bitcoin ETC (BTCE.DE)', 'Criptoactivo volátil'],
    ]
)

doc.add_heading('Máximo Sharpe (orientativo)', level=2)
add_table(
    ['Activo', 'Peso (%)'],
    [
        ['S&P 500', '30'],
        ['Info Tech Sector', '25'],
        ['World Small Cap', '20'],
        ['Emerging Markets', '15'],
        ['Bitcoin ETC', '10'],
    ]
)

doc.add_heading('Mínima Volatilidad (orientativo)', level=2)
add_table(
    ['Activo', 'Peso (%)'],
    [
        ['S&P 500', '40'],
        ['Info Tech Sector', '15'],
        ['World Small Cap', '15'],
        ['Emerging Markets', '20'],
        ['Bitcoin ETC', '10'],
    ]
)

doc.add_heading('3. Estadísticas Reales por Activo', level=1)
add_table(
    ['Activo', 'Ret. anual (%)', 'Vol. anual (%)', 'Max DD (%)', 'Sharpe'],
    activos_rows
)

doc.add_heading('4. Matriz de Correlación', level=1)
short = [a[:15] for a in nombres_ok]
corr_headers = [''] + short
corr_data = []
for i, t in enumerate(activos_ok):
    row = [short[i]] + [f"{corr_named.iloc[i,j]:.2f}" for j in range(len(activos_ok))]
    corr_data.append(row)
add_table(corr_headers, corr_data)

doc.add_heading('5. Carteras Óptimas Reales (Datos de Hoy)', level=1)
doc.add_heading('5.1 Máximo Sharpe', level=2)
doc.add_paragraph(
    f'Retorno: {res_s[0]*100:.1f}% | Volatilidad: {res_s[1]*100:.1f}% | Sharpe: {res_s[2]:.3f}'
)
for k, v in pesos_s.items():
    doc.add_paragraph(f'{k}: {v:.1f}%', style='List Bullet')

doc.add_heading('5.2 Mínima Volatilidad', level=2)
doc.add_paragraph(
    f'Retorno: {res_v[0]*100:.1f}% | Volatilidad: {res_v[1]*100:.1f}% | Sharpe: {res_v[2]:.3f}'
)
for k, v in pesos_v.items():
    doc.add_paragraph(f'{k}: {v:.1f}%', style='List Bullet')

doc.add_heading('6. Recomendaciones', level=1)
doc.add_paragraph(
    'La cartera de Máximo Sharpe ofrece la mejor relación riesgo/retorno incluso en perfiles agresivos. '
    'La de Mínima Volatilidad, aunque con menor volatilidad, sigue siendo arriesgada. '
    'El Bitcoin puede aumentar drásticamente tanto la rentabilidad como la volatilidad. '
    'Se recomienda rebalanceo semestral y solo para inversores con alta tolerancia al riesgo.'
)

dis = doc.add_paragraph()
dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dis.add_run('Disclaimer: Este informe no es asesoramiento financiero. Rentabilidades pasadas no garantizan resultados futuros.')
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save('Informe_Arriesgado.docx')
print("✅ Documento Word generado: Informe_Arriesgado.docx")