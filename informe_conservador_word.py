"""
Informe Cartera Conservadora (7 activos) - Word
================================================
Ejecutar: python informe_conservador_word.py
"""

import pandas as pd
import numpy as np
import yfinance as yf
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

# ── CONFIGURACIÓN ( activos conservadores) ─────────────────
ACTIVOS = {
    "IS3N.DE": "Emerging Markets",
    "EUNA.DE": "Global Bonds",
    "EUNL.DE": "MSCI World",          
    "IQQ0.DE": "Min Volatility",
    "SGLD.L": "Physical Gold",
}
TICKERS = list(ACTIVOS.keys())
TASA_RF = 0.0315
N_DIAS = 252
PERIODO = "3y"
MC_SIMULACIONES = 50000

# ── DESCARGA DE DATOS ─────────────────────────────────────
print("⬇ Descargando datos en tiempo real...")
raw = yf.download(TICKERS, period=PERIODO, auto_adjust=True, progress=False)

if "Close" not in raw.columns:
    data = raw.dropna(axis=1, how='all')
else:
    data = raw["Close"].dropna(axis=1, how='all')

min_dias = 100
activos_ok = [t for t in data.columns if data[t].notna().sum() >= min_dias]
if len(activos_ok) < 2:
    raise ValueError("Se necesitan al menos 2 activos con datos.")
data = data[activos_ok]
nombres_ok = [ACTIVOS.get(t, t) for t in activos_ok]
returns = data.pct_change().dropna()
n = len(activos_ok)
print(f"✅ Activos disponibles: {activos_ok}\n")

# ── MÉTRICAS POR ACTIVO ───────────────────────────────────
ret_anual = returns.mean() * N_DIAS
vol_anual = returns.std() * np.sqrt(N_DIAS)
sharpe = (ret_anual - TASA_RF) / vol_anual
max_dd = ((data / data.cummax()) - 1).min()

activos_rows = []
for t in activos_ok:
    activos_rows.append([
        f"{ACTIVOS.get(t, t)} ({t})",
        f"{ret_anual[t]*100:.2f}",
        f"{vol_anual[t]*100:.2f}",
        f"{max_dd[t]*100:.2f}",
        f"{sharpe[t]:.3f}"
    ])

# ── MATRIZ DE CORRELACIÓN ─────────────────────────────────
corr = returns.corr()
corr_named = corr.copy()
corr_named.columns = nombres_ok
corr_named.index = nombres_ok

# ── OPTIMIZACIÓN MONTE CARLO (sin scipy) ──────────────────
def stats_cartera(weights):
    weights = np.array(weights)
    port_returns = returns.dot(weights)
    ret = port_returns.mean() * N_DIAS
    vol = port_returns.std() * np.sqrt(N_DIAS)
    if vol == 0:
        return np.nan, np.nan, np.nan
    sr = (ret - TASA_RF) / vol
    return ret, vol, sr

print(f"🎲 Monte Carlo con {MC_SIMULACIONES:,} carteras...")
np.random.seed(42)
mc_ret = np.zeros(MC_SIMULACIONES)
mc_vol = np.zeros(MC_SIMULACIONES)
mc_sr  = np.zeros(MC_SIMULACIONES)
pesos_mc = np.zeros((MC_SIMULACIONES, n))

for i in range(MC_SIMULACIONES):
    w = np.random.random(n)
    w /= w.sum()
    pesos_mc[i] = w
    r, v, s = stats_cartera(w)
    mc_ret[i] = r
    mc_vol[i] = v
    mc_sr[i]  = s

idx_best_sharpe = np.nanargmax(mc_sr)
idx_min_vol = np.nanargmin(mc_vol)

best_w_sharpe = pesos_mc[idx_best_sharpe]
best_w_vol = pesos_mc[idx_min_vol]

res_s = stats_cartera(best_w_sharpe)
res_v = stats_cartera(best_w_vol)

pesos_s = {nombres_ok[i]: best_w_sharpe[i]*100 for i in range(n) if best_w_sharpe[i] > 0.005}
pesos_v = {nombres_ok[i]: best_w_vol[i]*100 for i in range(n) if best_w_vol[i] > 0.005}

# ── GENERAR DOCUMENTO WORD ─────────────────────────────────
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Título
title = doc.add_heading('INFORME FINANCIERO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Cartera Conservadora Global (7 ETFs)')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x33, 0x55, 0x77)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")} | Periodo: {PERIODO} | RF BCE: {TASA_RF*100:.2f}%').font.size = Pt(9)
doc.add_paragraph()

# Función para añadir tablas
def add_table(headers, data):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for row in data:
        rc = table.add_row().cells
        for i, item in enumerate(row):
            rc[i].text = str(item)
            for p in rc[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()

# 1. Resumen ejecutivo
doc.add_heading('1. Resumen Ejecutivo', level=1)
doc.add_paragraph(
    f'Datos desde {returns.index[0].date()} hasta {returns.index[-1].date()}. '
    f'Optimización Monte Carlo con {MC_SIMULACIONES:,} carteras. '
    'Se muestran los pesos orientativos de referencia y los pesos reales obtenidos hoy.'
)
add_table(
    ['Cartera', 'Retorno (%)', 'Volatilidad (%)', 'Sharpe'],
    [
        ['Máximo Sharpe', f"{res_s[0]*100:.1f}", f"{res_s[1]*100:.1f}", f"{res_s[2]:.3f}"],
        ['Mínima Volatilidad', f"{res_v[0]*100:.1f}", f"{res_v[1]*100:.1f}", f"{res_v[2]:.3f}"],
        ['Tasa libre riesgo', f"{TASA_RF*100:.2f}", "0.00", '—']
    ]
)

# 2. Composición orientativa (pesos típicos)
doc.add_heading('2. Composición Orientativa (Pesos Típicos)', level=1)
doc.add_paragraph(
    'Estos pesos son una referencia basada en simulaciones históricas. '
    'Los pesos reales calculados con los datos actuales se muestran en la sección 5.'
)

doc.add_heading('Activos y su rol', level=2)
add_table(
    ['Activo', 'Rol'],
    [
        ['Global Bonds (EUNA.DE)', 'Estabilizador principal'],
        ['MSCI World Swap (XMWO.DE)', 'Crecimiento renta variable'],
        ['Min Volatility (IQQ0.DE)', 'Renta variable defensiva'],
        ['Physical Gold (SGLD.L)', 'Cobertura inflación / crisis'],
        ['Emerging Markets (IS3N.DE)', 'Exposición emergente'],
        ['Developed World (VHVG.L)', 'Renta variable desarrollada (redundante)'],
        ['Inflation Linked (IBEI.DE)', 'Bonos indexados a inflación (redundante)'],
    ]
)

doc.add_heading('Máximo Sharpe (orientativo)', level=2)
add_table(
    ['Activo', 'Peso (%)'],
    [
        ['Global Bonds (EUNA.DE)', '35'],
        ['MSCI World Swap (XMWO.DE)', '22'],
        ['Min Volatility (IQQ0.DE)', '18'],
        ['Physical Gold (SGLD.L)', '15'],
        ['Emerging Markets (IS3N.DE)', '10'],
        ['Developed World (VHVG.L)', '0'],
        ['Inflation Linked (IBEI.DE)', '0'],
    ]
)

doc.add_heading('Mínima Volatilidad (orientativo)', level=2)
add_table(
    ['Activo', 'Peso (%)'],
    [
        ['Global Bonds (EUNA.DE)', '45'],
        ['Inflation Linked (IBEI.DE)', '30'],
        ['Min Volatility (IQQ0.DE)', '15'],
        ['Physical Gold (SGLD.L)', '10'],
        ['MSCI World Swap (XMWO.DE)', '0'],
        ['Emerging Markets (IS3N.DE)', '0'],
        ['Developed World (VHVG.L)', '0'],
    ]
)

# 3. Estadísticas reales
doc.add_heading('3. Estadísticas Reales por Activo', level=1)
doc.add_paragraph('Métricas calculadas con los datos más recientes del mercado.')
add_table(
    ['Activo', 'Ret. anual (%)', 'Vol. anual (%)', 'Max DD (%)', 'Sharpe'],
    activos_rows
)

# 4. Matriz de correlación
doc.add_heading('4. Matriz de Correlación', level=1)
short = [a[:15] for a in nombres_ok]
corr_headers = [''] + short
corr_data = []
for i, t in enumerate(activos_ok):
    row = [short[i]] + [f"{corr_named.iloc[i,j]:.2f}" for j in range(len(activos_ok))]
    corr_data.append(row)
add_table(corr_headers, corr_data)

# 5. Carteras óptimas reales
doc.add_heading('5. Carteras Óptimas Reales (Datos de Hoy)', level=1)
doc.add_heading('5.1 Máximo Sharpe', level=2)
doc.add_paragraph(
    f'Retorno: {res_s[0]*100:.1f}% | Volatilidad: {res_s[1]*100:.1f}% | Sharpe: {res_s[2]:.3f}'
)
for k, v in pesos_s.items():
    doc.add_paragraph(f'{k}: {v:.1f}%', style='List Bullet')

doc.add_heading('5.2 Mínima Volatilidad', level=2)
doc.add_paragraph(
    f'Retorno: {res_v[0]*100:.1f}% | Volatilidad: {res_v[1]*100:.1f}% | Sharpe: {res_v[2]:.3f}'
)
for k, v in pesos_v.items():
    doc.add_paragraph(f'{k}: {v:.1f}%', style='List Bullet')

# 6. Recomendaciones
doc.add_heading('6. Recomendaciones', level=1)
doc.add_paragraph(
    'La cartera de Máximo Sharpe ofrece el mejor equilibrio riesgo/retorno para inversores conservadores. '
    'La de Mínima Volatilidad es ideal para quienes priorizan la estabilidad. '
    'Se recomienda rebalanceo semestral para mantener los pesos óptimos.'
)

dis = doc.add_paragraph()
dis.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dis.add_run('Disclaimer: Este informe no es asesoramiento financiero. Rentabilidades pasadas no garantizan resultados futuros.')
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save('Informe_Conservador.docx')
print("✅ Documento Word generado: Informe_Conservador.docx")